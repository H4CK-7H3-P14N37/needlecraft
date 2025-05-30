import lxml
from docx import Document
from docx.oxml import ns
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.table import _Cell
import csv
# from docx.oxml import OxmlElement


class OSSReportGen:
    def __init__(self, save_dir="/tmp/"):
        self.COMPANY_LOGO = "static/img/redteam.png"
        self.RISK_LEVEL_CRITICAL_RGB = RGBColor(152, 0, 0)
        self.RISK_LEVEL_CRITICAL_HEX = '980000'
        self.RISK_LEVEL_HIGH_RGB = RGBColor(255, 0, 0)
        self.RISK_LEVEL_HIGH_HEX = 'ff0000'
        self.RISK_LEVEL_MODERATE_RGB = RGBColor(241, 194, 50)
        self.RISK_LEVEL_MODERATE_HEX = 'f1c232'
        self.RISK_LEVEL_LOW_RGB = RGBColor(145, 219, 16)
        self.RISK_LEVEL_LOW_HEX = '91db10'
        self.RISK_LEVEL_INFO_RGB = RGBColor(3, 187, 72)
        self.RISK_LEVEL_INFO_HEX = '03bb48'
        self.save_dir = save_dir
        if not self.save_dir.endswith("/"):
            self.save_dir = f"{self.save_dir}/"

    def format_vuln_data_file(self, csv_input_file_location):
        vuln_list = []
        with open(csv_input_file_location, newline='') as csvfile:
            vuln_reader = csv.DictReader(csvfile)
            vuln_list = [row for row in vuln_reader]
        return vuln_list

    def format_vuln_description_text(self, vuln_description):
        end_text = ""
        if vuln_description:
            desc_text_list = vuln_description.strip().split("\n")
            if desc_text_list:
                for current_text in desc_text_list:
                    if current_text:
                        end_text += " {}".format(
                            current_text
                        )
                    else:
                        end_text += "\n\n"
        return end_text

    def format_vuln_data_to_dict(
            self,
            valid_vuln_list,
            filter_to_scores=False,
            score_override={},
            solutions_override_dict={}):
        vuln_dict = {}
        # override score here based on input
        if score_override:
            for vuln_dict_tmp in valid_vuln_list:
                if vuln_dict_tmp.get("Name") in score_override.keys():
                    if 'CVSS_v3_0_Base_Score' in vuln_dict_tmp.keys():
                        vuln_dict_tmp.update(
                            {'CVSS_v3_0_Base_Score': score_override.get(vuln_dict_tmp.get("Name"))})
                    else:
                        vuln_dict_tmp.update(
                            {'CVSS v3.0 Base Score': score_override.get(vuln_dict_tmp.get("Name"))})
                    if 'CVSS_v2_0_Base_Score' in vuln_dict_tmp.keys():
                        vuln_dict_tmp.update(
                            {'CVSS_v2_0_Base_Score': score_override.get(vuln_dict_tmp.get("Name"))})
                    else:
                        vuln_dict_tmp.update(
                            {'CVSS v2.0 Base Score': score_override.get(vuln_dict_tmp.get("Name"))})

        if filter_to_scores:
            # CVSS 3.0 Base Score Range
            # Informational    0.0
            # Low              0.1-3.9
            # Medium           4.0-6.9
            # High             7.0-8.9
            # Critical         9.0-10.0
            tmp_valid_vuln_list = []
            for vuln in valid_vuln_list:
                use_cvss2 = False
                if 'CVSS_v3_0_Base_Score' in vuln.keys() or 'CVSS v3.0 Base Score' in vuln.keys():
                    if vuln.get('CVSS_v3_0_Base_Score') and vuln.get(
                            'CVSS_v3_0_Base_Score') != "None":
                        tmp_valid_vuln_list.append(vuln)
                    elif vuln.get('CVSS v3.0 Base Score') and vuln.get('CVSS v3.0 Base Score') != "None":
                        tmp_valid_vuln_list.append(vuln)
                    else:
                        use_cvss2 = True
                if use_cvss2:
                    if 'CVSS_v2_0_Base_Score' in vuln.keys() or 'CVSS v2.0 Base Score' in vuln.keys():
                        if vuln.get('CVSS_v2_0_Base_Score') and vuln.get(
                                'CVSS_v2_0_Base_Score') != "None":
                            tmp_valid_vuln_list.append(vuln)
                        if vuln.get('CVSS v2.0 Base Score') and vuln.get(
                                'CVSS v2.0 Base Score') != "None":
                            tmp_valid_vuln_list.append(vuln)
            valid_vuln_list = tmp_valid_vuln_list
        # discard garbage items
        pci_list_we_keep = []
        for vuln_pci_dict in valid_vuln_list:
            if vuln_pci_dict.get("Name").lower().strip(
            ) == "pci dss compliance : insecure communication has been detected":
                if 'CVSS_v3_0_Base_Score' in vuln_pci_dict.keys():
                    vuln_pci_dict.update(
                        {"Name": "Insecure Communication Has Been Detected", "CVSS_v3_0_Base_Score": "10"})
                else:
                    vuln_pci_dict.update(
                        {"Name": "Insecure Communication Has Been Detected", "CVSS v3.0 Base Score": "10"})
                pci_list_we_keep.append(vuln_pci_dict)

        valid_vuln_list = [vuln for vuln in valid_vuln_list if vuln.get(
            "Name") != "Internet Key Exchange (IKE) Aggressive Mode with Pre-Shared Key" and not vuln.get("Name").lower().startswith("pci dss compliance")]
        if pci_list_we_keep:
            valid_vuln_list.extend(pci_list_we_keep)
        for valid_vuln in valid_vuln_list:
            # we use the CVSS v3 score if it's there, but if not, pull the
            # CVSS v2 score and convert it to CVSS v3.
            use_cvss2 = False
            if 'CVSS_v3_0_Base_Score' in valid_vuln.keys():
                vuln_risk_score = valid_vuln.get('CVSS_v3_0_Base_Score')
                if not vuln_risk_score:
                    use_cvss2 = True
            else:
                vuln_risk_score = valid_vuln.get('CVSS v2.0 Base Score')
                if not vuln_risk_score:
                    use_cvss2 = True
            if use_cvss2:
                if 'CVSS_v2_0_Base_Score' in valid_vuln.keys():
                    vuln_risk_score = valid_vuln.get('CVSS_v2_0_Base_Score')
                else:
                    vuln_risk_score = valid_vuln.get('CVSS v2.0 Base Score')
            vuln_risk = "informational"
            if vuln_risk_score and vuln_risk_score != "None":
                if 0 < float(vuln_risk_score) < 4.0:
                    vuln_risk = "low"
                elif 3.9 < float(vuln_risk_score) < 7.0:
                    vuln_risk = "medium"
                elif 6.9 < float(vuln_risk_score) < 9.0:
                    vuln_risk = "high"
                elif 8.9 < float(vuln_risk_score) <= 10.0:
                    vuln_risk = "critical"
            valid_vuln.update({"Risk": vuln_risk.title()})
            vuln_name = valid_vuln.get('Name')
            if "(PCI DSS)" in vuln_name or "(PCI-DSS check)" in vuln_name:
                vuln_name = vuln_name.replace(
                    "(PCI-DSS check)",
                    "").replace(
                    "(PCI DSS)",
                    "").strip()
            vuln_proto = valid_vuln.get('Protocol')
            vuln_host = valid_vuln.get('Host')
            vuln_port = valid_vuln.get('Port')
            vuln_host_list = vuln_host.splitlines()
            vuln_proto_host_port = ""
            if vuln_proto and vuln_port:
                for vuln_host_address in vuln_host_list:
                    vuln_proto_host_port += "{} {}:{}\n".format(
                        vuln_proto, vuln_host_address, vuln_port)
            else:
                for vuln_host_address in vuln_host_list:
                    vuln_proto_host_port += "{}\n".format(vuln_host)
            vuln_proto_host_port = vuln_proto_host_port.rstrip('\n')
            vuln_desc = valid_vuln.get('Description')
            vuln_risk_desc = valid_vuln.get('Synopsis')
            vuln_solution = valid_vuln.get('Solution')
            org_vuln_name = valid_vuln.get('Name')
            if 'See_Also' in valid_vuln.keys():
                vuln_extra_resources = valid_vuln.get('See_Also')
            else:
                vuln_extra_resources = valid_vuln.get('See Also')
            vuln_solution_screenshot_list = []
            if org_vuln_name:
                solution_override_dict = solutions_override_dict.get(
                    org_vuln_name)
                if solution_override_dict:
                    vuln_solution = solution_override_dict.get('solution_text')
                    vuln_solution_screenshot_list = solution_override_dict.get(
                        'screenshots')
                    vuln_extra_resources = solution_override_dict.get(
                        'solution_see_also')

            vuln_screenshots = valid_vuln.get("finding_screenshots")
            if 'Plugin_Output' in valid_vuln.keys():
                vuln_evidence = valid_vuln.get('Plugin_Output')
            else:
                vuln_evidence = valid_vuln.get('Plugin Output')
            cve_number = valid_vuln.get('CVE')
            if vuln_risk not in vuln_dict.keys():
                vuln_dict.update({vuln_risk: {}})
            if vuln_name not in vuln_dict[vuln_risk].keys():
                vuln_dict[vuln_risk].update({vuln_name: {}})
            if "description" not in vuln_dict[vuln_risk][vuln_name].keys():
                vuln_dict[vuln_risk][vuln_name]["description"] = self.format_vuln_description_text(
                    vuln_desc)
            if "impact" not in vuln_dict[vuln_risk][vuln_name].keys():
                vuln_dict[vuln_risk][vuln_name]["impact"] = self.format_vuln_description_text(
                    vuln_risk_desc)
            if "recommendations" not in vuln_dict[vuln_risk][vuln_name].keys():
                vuln_dict[vuln_risk][vuln_name]["recommendations"] = self.format_vuln_description_text(
                    vuln_solution)
                vuln_dict[vuln_risk][vuln_name]["recommendations_screenshots"] = vuln_solution_screenshot_list
            if "evidence" not in vuln_dict[vuln_risk][vuln_name].keys():
                evidence_response = self.format_vuln_description_text(
                    vuln_evidence)
                if evidence_response.strip():
                    vuln_dict[vuln_risk][vuln_name]["evidence"] = [
                        evidence_response]
            else:
                evidence_response = self.format_vuln_description_text(
                    vuln_evidence)
                if evidence_response.strip():
                    vuln_dict[vuln_risk][vuln_name]["evidence"].append(
                        evidence_response)
            if "hosts_list" not in vuln_dict[vuln_risk][vuln_name].keys():
                vuln_dict[vuln_risk][vuln_name]["hosts_list"] = []
            if "cve_list" not in vuln_dict[vuln_risk][vuln_name].keys():
                vuln_dict[vuln_risk][vuln_name]["cve_list"] = []
            if "sources_list" not in vuln_dict[vuln_risk][vuln_name].keys():
                vuln_dict[vuln_risk][vuln_name]["sources_list"] = []
            if vuln_proto_host_port:
                if vuln_proto_host_port not in vuln_dict[vuln_risk][vuln_name]["hosts_list"]:
                    if vuln_proto_host_port.endswith(":0"):
                        vuln_proto_host_port = vuln_proto_host_port[:-2]
                    if ',' in vuln_proto_host_port:
                        vuln_proto_host_port = vuln_proto_host_port.split(",")
                    elif '\n' in vuln_proto_host_port:
                        vuln_proto_host_port = vuln_proto_host_port.split("\n")
                    else:
                        vuln_proto_host_port = [vuln_proto_host_port]
                    vuln_dict[vuln_risk][vuln_name]["hosts_list"].extend(
                        vuln_proto_host_port)
            if cve_number:
                if ',' in cve_number:
                    cve_number = cve_number.split(",")
                elif '\n' in cve_number:
                    cve_number = cve_number.split("\n")
                else:
                    cve_number = [cve_number]
                for cve_num in cve_number:
                    if cve_num not in vuln_dict[vuln_risk][vuln_name]["cve_list"]:
                        vuln_dict[vuln_risk][vuln_name]["cve_list"].append(
                            cve_num.lower())
            if vuln_extra_resources:
                if ',' in vuln_extra_resources:
                    vuln_extra_resources = vuln_extra_resources.split(",")
                elif '\n' in vuln_extra_resources:
                    vuln_extra_resources = vuln_extra_resources.split("\n")
                else:
                    vuln_extra_resources = [vuln_extra_resources]
                for vuln_ex_resc in vuln_extra_resources:
                    if vuln_ex_resc not in vuln_dict[vuln_risk][vuln_name]["sources_list"]:
                        vuln_dict[vuln_risk][vuln_name]["sources_list"].append(
                            vuln_ex_resc)
            if vuln_screenshots:
                if "finding_screenshots" in vuln_dict[vuln_risk][vuln_name].keys(
                ):
                    vuln_dict[vuln_risk][vuln_name]["finding_screenshots"].extend(
                        vuln_screenshots)
                else:
                    vuln_dict[vuln_risk][vuln_name]["finding_screenshots"] = vuln_screenshots
            else:
                if "finding_screenshots" not in vuln_dict[vuln_risk][vuln_name].keys(
                ):
                    vuln_dict[vuln_risk][vuln_name]["finding_screenshots"] = []
        return vuln_dict

    def risk_level_color(self, risk_level):
        if risk_level.lower() == "informational":
            return self.RISK_LEVEL_INFO_RGB, self.RISK_LEVEL_INFO_HEX
        elif risk_level.lower() == "low":
            return self.RISK_LEVEL_LOW_RGB, self.RISK_LEVEL_LOW_HEX
        elif risk_level.lower() == "moderate":
            return self.RISK_LEVEL_MODERATE_RGB, self.RISK_LEVEL_MODERATE_HEX
        elif risk_level.lower() == "medium":
            return self.RISK_LEVEL_MODERATE_RGB, self.RISK_LEVEL_MODERATE_HEX
        elif risk_level.lower() == "high":
            return self.RISK_LEVEL_HIGH_RGB, self.RISK_LEVEL_HIGH_HEX
        elif risk_level.lower() == "critical":
            return self.RISK_LEVEL_CRITICAL_RGB, self.RISK_LEVEL_CRITICAL_HEX
        return self.RISK_LEVEL_INFO_RGB, self.RISK_LEVEL_INFO_HEX

    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(self, run):
        fldChar1 = self.create_element('w:fldChar')
        self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
        instrText = self.create_element('w:instrText')
        self.create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"
        fldChar2 = self.create_element('w:fldChar')
        self.create_attribute(fldChar2, 'w:fldCharType', 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def set_updatefields_true(self, docx_path):
        namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        doc = Document(docx_path)
        # add child to doc.settings element
        element_updatefields = lxml.etree.SubElement(
            doc.settings.element, f"{namespace}updateFields"
        )
        element_updatefields.set(f"{namespace}val", "true")
        doc.save(docx_path)  # Heading ##

    def add_line_breaks(self, doc, n=1):
        for _ in range(n):
            blank_paragraph = doc.add_paragraph(" ")
            blank_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    def indent_table(self, table, indent):
        # noinspection PyProtectedMember
        tbl_pr = table._element.xpath('w:tblPr')
        if tbl_pr:
            e = OxmlElement('w:tblInd')
            e.set(qn('w:w'), str(indent))
            e.set(qn('w:type'), 'dxa')
            tbl_pr[0].append(e)

    def set_cell_border(self, cell: _Cell, **kwargs):
        """
        Set cell`s border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 12, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)),
                                    str(edge_data[key]))

    def set_cell_borders_clear(self, cell):
        self.set_cell_border(
            cell,
            top={
                "sz": 0,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 10,
                "insideV": 10},
            bottom={
                "sz": 0,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 10,
                "insideV": 10},
            start={
                "sz": 0,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 10,
                "insideV": 10},
            end={
                "sz": 0,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 10,
                "insideV": 10},
        )

    def set_cell_borders_clear_1(self, cell):
        self.set_cell_border(
            cell,
            top={
                "sz": 12,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
            bottom={
                "sz": 12,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
            start={
                "sz": 12,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
            end={
                "sz": 12,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
        )

    def set_cell_borders_black(self, cell):
        self.set_cell_border(
            cell,
            top={
                "sz": 12,
                "val": "single",
                "color": "#000000",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
            bottom={
                "sz": 12,
                "val": "single",
                "color": "#000000",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
            start={
                "sz": 12,
                "val": "single",
                "color": "#000000",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
            end={
                "sz": 12,
                "val": "single",
                "color": "#000000",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
        )

    def set_cell_borders_black_bottom(self, cell):
        self.set_cell_border(
            cell,
            top={
                "sz": 0,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
            bottom={
                "sz": 12,
                "val": "single",
                "color": "#000000",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
            start={
                "sz": 0,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
            end={
                "sz": 0,
                "val": "single",
                "color": "#FFFFFF",
                "space": "0",
                "insideH": 0,
                "insideV": 0},
        )

    def set_cell_margins(self, cell: _Cell, **kwargs):
        """
        cell:  actual cell instance you want to modify

        usage:

            set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

        provided values are in twentieths of a point (1/1440 of an inch).
        read more here: http://officeopenxml.com/WPtableCellMargins.php
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for m in [
            "top",
            "start",
            "bottom",
            "end",
        ]:
            if m in kwargs:
                node = OxmlElement("w:{}".format(m))
                node.set(qn('w:w'), str(kwargs.get(m)))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

        tcPr.append(tcMar)

    def set_assessment_page_margins(self, cell):
        self.set_cell_margins(cell, top=10, start=10, bottom=10, end=10)

    def set_findings_table_margins(self, cell):
        self.set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    def get_friendly_test_name(self, test_type):
        friendly_test_type_name = ""
        test_types = [
            "vuln_scan",
            "vuln_assessment",
            "pentest",
            "webapp",
        ]
        test_type = test_type.lower()
        if test_type in test_types:
            if test_type == "vuln_scan":
                friendly_test_type_name = "vulnerability scan"
            elif test_type == "vuln_assessment":
                friendly_test_type_name = "vulnerability assessment"
            elif test_type == "pentest":
                friendly_test_type_name = "penetration test"
            elif test_type == "webapp":
                friendly_test_type_name = "web application penetration test"
        return friendly_test_type_name

    def get_friendly_scope_test_from(self, scope_test_from):
        friendly_scope_test_from = ""
        scope_test_froms = [
            "internal",
            "external",
            "both"
        ]
        scope_test_from = scope_test_from.lower()
        if scope_test_from in scope_test_froms:
            if scope_test_from == "internal":
                friendly_scope_test_from = "internal infrastructure"
            elif scope_test_from == "external":
                friendly_scope_test_from = "external infrastructure"
            elif scope_test_from == "both":
                friendly_scope_test_from = "internal and external infrastructure"
        return friendly_scope_test_from

    def get_compromise_status(self, compromise_success):
        if compromise_success:
            return "successful"
        return "unsuccessful"

    def get_risk_color_codes(self, risk_level):
        risk_level = risk_level.lower()
        risk_level_color_rgb = self.RISK_LEVEL_INFO_RGB
        risk_level_color_hex = self.RISK_LEVEL_INFO_HEX
        if risk_level == "very low":
            risk_level_color_rgb = self.RISK_LEVEL_INFO_RGB
            risk_level_color_hex = self.RISK_LEVEL_INFO_HEX
        elif risk_level == "low":
            risk_level_color_rgb = self.RISK_LEVEL_LOW_RGB
            risk_level_color_hex = self.RISK_LEVEL_LOW_HEX
        elif risk_level == "moderate":
            risk_level_color_rgb = self.RISK_LEVEL_MODERATE_RGB
            risk_level_color_hex = self.RISK_LEVEL_MODERATE_HEX
        elif risk_level == "medium":
            risk_level_color_rgb = self.RISK_LEVEL_MODERATE_RGB
            risk_level_color_hex = self.RISK_LEVEL_MODERATE_HEX
        elif risk_level == "high":
            risk_level_color_rgb = self.RISK_LEVEL_HIGH_RGB
            risk_level_color_hex = self.RISK_LEVEL_HIGH_HEX
        elif risk_level == "critical":
            risk_level_color_rgb = self.RISK_LEVEL_CRITICAL_RGB
            risk_level_color_hex = self.RISK_LEVEL_CRITICAL_HEX
        return risk_level_color_rgb, risk_level_color_hex

    def format_vuln_data_into_dict(self,
                                   vuln_scan_csv_file_path="",
                                   vuln_list=[],
                                   filter_to_scores=True,
                                   score_overrides_dict={},
                                   solutions_override_dict={}
                                   ):
        vulnerability_dict = {}
        if vuln_scan_csv_file_path:
            vulnerability_list = self.format_vuln_data_file(
                vuln_scan_csv_file_path)
            vulnerability_dict = self.format_vuln_data_to_dict(
                vulnerability_list,
                filter_to_scores=filter_to_scores,
                score_override=score_overrides_dict
            )
        elif vuln_list:
            vulnerability_dict = self.format_vuln_data_to_dict(
                vuln_list,
                filter_to_scores=filter_to_scores,
                score_override=score_overrides_dict,
                solutions_override_dict=solutions_override_dict
            )
        else:
            raise Exception("no vulnerability list given")
        return vulnerability_dict

    def setup_document_and_styles(self):
        # create the document
        self.document = Document()

        # Setup styles
        RT_PRIMARY_COLOR_R = 255
        RT_PRIMARY_COLOR_G = 75
        RT_PRIMARY_COLOR_B = 43
        green_header1 = self.document.styles.add_style(
            'rt_header', WD_STYLE_TYPE.PARAGRAPH)
        green_header1.font.name = "Montserrat"
        green_header1.font.size = Pt(48)
        green_header1.font.color.rgb = RGBColor(
            RT_PRIMARY_COLOR_R, RT_PRIMARY_COLOR_G, RT_PRIMARY_COLOR_B)

        green_header2 = self.document.styles.add_style(
            'rt_header_sub', WD_STYLE_TYPE.PARAGRAPH)
        green_header2.font.name = "Montserrat"
        green_header2.font.size = Pt(30)
        green_header2.font.color.rgb = RGBColor(
            RT_PRIMARY_COLOR_R, RT_PRIMARY_COLOR_G, RT_PRIMARY_COLOR_B)

        green_page_header = self.document.styles.add_style(
            'rt_page_green', WD_STYLE_TYPE.CHARACTER)
        green_page_header.font.name = "Montserrat"
        green_page_header.font.size = Pt(9)
        green_page_header.font.color.rgb = RGBColor(
            RT_PRIMARY_COLOR_R, RT_PRIMARY_COLOR_G, RT_PRIMARY_COLOR_B)

        black_page_header = self.document.styles.add_style(
            'rt_page_black', WD_STYLE_TYPE.CHARACTER)
        black_page_header.font.name = "Montserrat"
        black_page_header.font.size = Pt(9)
        black_page_header.font.color.rgb = RGBColor(0, 0, 0)

        page_header_one = self.document.styles['Heading 1']
        page_header_one.font.name = "Lato"
        page_header_one.font.size = Pt(18)
        page_header_one.font.color.rgb = RGBColor(
            RT_PRIMARY_COLOR_R, RT_PRIMARY_COLOR_G, RT_PRIMARY_COLOR_B)

        page_header_two = self.document.styles['Heading 2']
        page_header_two.font.name = "Lato"
        page_header_two.font.size = Pt(16)
        page_header_two.font.color.rgb = RGBColor(0, 0, 0)

        page_header_three = self.document.styles['Heading 3']
        page_header_three.font.name = "Lato"
        page_header_three.font.size = Pt(13)
        page_header_three.font.color.rgb = RGBColor(0, 0, 0)

        body_text_black = self.document.styles.add_style(
            'rt_text_black', WD_STYLE_TYPE.PARAGRAPH)
        body_text_black.font.name = "Lato"
        body_text_black.font.size = Pt(11)
        body_text_black.font.color.rgb = RGBColor(0, 0, 0)

        header_text_black = self.document.styles.add_style(
            'rt_header_text_black', WD_STYLE_TYPE.PARAGRAPH)
        header_text_black.font.name = "Lato"
        header_text_black.font.size = Pt(13)
        header_text_black.font.color.rgb = RGBColor(0, 0, 0)

        rt_text_black_subtext = self.document.styles.add_style(
            'rt_text_black_subtext', WD_STYLE_TYPE.PARAGRAPH)
        rt_text_black_subtext.font.name = "Lato"
        rt_text_black_subtext.font.size = Pt(9)
        rt_text_black_subtext.font.color.rgb = RGBColor(0, 0, 0)

        rt_text_white_subtext = self.document.styles.add_style(
            'rt_text_white_subtext', WD_STYLE_TYPE.PARAGRAPH)
        rt_text_white_subtext.font.name = "Lato"
        rt_text_white_subtext.font.size = Pt(9)
        rt_text_white_subtext.font.color.rgb = RGBColor(240, 238, 236)

        rt_text_center_black_subtext = self.document.styles.add_style(
            'rt_text_center_black_subtext', WD_STYLE_TYPE.PARAGRAPH)
        rt_text_center_black_subtext.font.name = "Lato"
        rt_text_center_black_subtext.font.size = Pt(9)
        rt_text_center_black_subtext.font.color.rgb = RGBColor(0, 0, 0)

        body_text_green = self.document.styles.add_style(
            'body_text_green', WD_STYLE_TYPE.PARAGRAPH)
        body_text_green.font.name = "Lato"
        body_text_green.font.size = Pt(10)
        body_text_green.font.color.rgb = RGBColor(
            RT_PRIMARY_COLOR_R, RT_PRIMARY_COLOR_G, RT_PRIMARY_COLOR_B)

        body_text_green_title = self.document.styles.add_style(
            'body_text_green_title', WD_STYLE_TYPE.PARAGRAPH)
        body_text_green_title.font.name = "Lato"
        body_text_green_title.font.size = Pt(14)
        body_text_green_title.font.color.rgb = RGBColor(
            RT_PRIMARY_COLOR_R, RT_PRIMARY_COLOR_G, RT_PRIMARY_COLOR_B)

        findings_table_text_title = self.document.styles.add_style(
            'findings_table_text_title', WD_STYLE_TYPE.PARAGRAPH)
        findings_table_text_title.font.name = "Lato"
        findings_table_text_title.font.size = Pt(12)
        findings_table_text_title.font.color.rgb = RGBColor(0, 0, 0)
        findings_table_text_title.font.bold = True

        findings_table_text = self.document.styles.add_style(
            'findings_table_text', WD_STYLE_TYPE.CHARACTER)
        findings_table_text.font.name = "Lato"
        findings_table_text.font.size = Pt(10)
        findings_table_text.font.color.rgb = RGBColor(0, 0, 0)

        findings_table_text_alt = self.document.styles.add_style(
            'findings_table_text_alt', WD_STYLE_TYPE.CHARACTER)
        findings_table_text_alt.font.name = "Lato"
        findings_table_text_alt.font.size = Pt(10)

        green_bullets = self.document.styles['List Bullet']
        green_bullets.font.name = "Lato"
        green_bullets.font.size = Pt(10)
        green_bullets.font.color.rgb = RGBColor(
            RT_PRIMARY_COLOR_R, RT_PRIMARY_COLOR_G, RT_PRIMARY_COLOR_B)

        green_bullets = self.document.styles['List Bullet 2']
        green_bullets.font.name = "Lato"
        green_bullets.font.size = Pt(11)
        green_bullets.font.color.rgb = RGBColor(0, 0, 0)

        bullets_black = self.document.styles['List Number']
        bullets_black.font.name = "Lato"
        bullets_black.font.size = Pt(11)
        bullets_black.font.color.rgb = RGBColor(0, 0, 0)
        bullets_black.font.bold = True

        sub_bullets_black = self.document.styles['List Number 2']
        sub_bullets_black.font.name = "Lato"
        sub_bullets_black.font.size = Pt(11)
        sub_bullets_black.font.color.rgb = RGBColor(0, 0, 0)

        sub3_bullets_black = self.document.styles['List Number 3']
        sub3_bullets_black.font.name = "Lato"
        sub3_bullets_black.font.size = Pt(11)
        sub3_bullets_black.font.color.rgb = RGBColor(0, 0, 0)

    def create_title_page(
            self,
            reporting_team,
            report_for,
            friendly_test_type_name,
            title_month_year):
        # START PAGE ONE
        section = self.document.sections[0]
        # PAGE 1 Header
        header = section.header
        paragraph = header.paragraphs[0]
        logo_run = paragraph.add_run()
        logo_run.add_picture(self.COMPANY_LOGO, width=Inches(0.75))
        header_run = paragraph.add_run()
        header_run.text = "\t\tReport For: "
        header_run.style = "rt_page_green"
        header_run_company = paragraph.add_run()
        header_run_company.text = report_for
        header_run_company.style = "rt_page_black"
        # Company page
        self.document.add_paragraph("_" * 105)
        report_for_header = self.document.add_paragraph(
            report_for, style='rt_header')
        report_for_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Adding spacing
        self.add_line_breaks(self.document, 4)
        # Title page service name
        test_type_header = self.document.add_paragraph(
            friendly_test_type_name.title(), style='rt_header_sub')
        test_type_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Adding spacing
        self.add_line_breaks(self.document, 4)
        # adding in date of report
        test_date = self.document.add_paragraph(
            title_month_year, style='rt_header_sub')
        test_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # By line on title page
        test_by_line = self.document.add_paragraph(
            f"By: {reporting_team}", style='rt_header_sub')
        test_by_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # adding in the footer
        footer = section.footer
        paragraph = footer.paragraphs[0]
        footer_run = paragraph.add_run()
        footer_run.text = f"\t{reporting_team} | {
            report_for} {friendly_test_type_name.title()}\t"
        footer_run.style = "rt_page_black"
        footer_run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.add_page_number(footer_run)
        self.document.add_page_break()
        # END OF PAGE 1

    def create_table_of_contents(self):
        # START Table of Contents
        # Title for TOC
        toc_header = self.document.add_paragraph('Table of Contents')
        toc_header.style = "Heading 1"
        toc_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc_paragraph = self.document.add_paragraph()
        toc_paragraph.style = "rt_text_black"
        toc_run = toc_paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        # change 1-3 depending on heading levels you need
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        r_element = toc_run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = toc_paragraph._p
        self.document.add_page_break()
        # END OF TOC Page 2

    def create_assessment_contact(
        self,
        reporting_team,
        report_for,
        friendly_test_type_name,
        start_date,
        end_date,
        rt_contact_data,
        client_contact_data
    ):
        # START page 3 Assessment Info
        table = self.document.add_table(rows=1, cols=2, style="Table Grid")
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(3.5)
        table.rows[0].cells[0].width = Inches(3.5)
        table.columns[1].width = Inches(4.5)
        table.rows[0].cells[1].width = Inches(4.5)
        self.indent_table(table, "-1in")
        shading_elm_1 = parse_xml(
            r'<w:shd {} w:fill="1c1c1c"/>'.format(nsdecls('w')))
        table_cell_0 = table.rows[0].cells[0]
        self.set_cell_borders_clear(table_cell_0)
        # set_assessment_page_margins(table_cell_0)
        table_cell_1 = table.rows[0].cells[1]
        # set_cell_borders_clear(table_cell_1)
        self.set_cell_borders_clear_1(table_cell_1)
        table_cell_0._tc.get_or_add_tcPr().append(shading_elm_1)
        self.add_line_breaks(table_cell_0, 1)
        table_cell_0.add_paragraph(
            "    About Red Teaming",
            style='body_text_green_title')
        self.add_line_breaks(table_cell_0, 1)
        table_cell_0.add_paragraph(
            "Red teaming is a method where a team simulates real-world threats to identify weaknesses and improve strategies.",
            style='List Bullet')
        self.add_line_breaks(table_cell_0, 1)
        table_cell_0.add_paragraph(
            "In cybersecurity, it mimics attackers' tactics to test and strengthen an organization's defenses.",
            style='List Bullet')
        self.add_line_breaks(table_cell_0, 1)
        table_cell_0.add_paragraph(
            "Outside of security, red teaming evaluates plans and strategies to uncover risks and refine decisions.",
            style='List Bullet')
        self.add_line_breaks(table_cell_0, 1)
        table_cell_0.add_paragraph(
            "Its main benefit is helping organizations proactively fix vulnerabilities and improve resilience.",
            style='List Bullet')
        self.add_line_breaks(table_cell_0, 2)
        assessment_info_header = table_cell_1.add_paragraph(
            '    Assessment Information', style='Heading 1')
        assessment_info_header.paragraph_format.space_before = Pt(0)
        assessment_info_header.paragraph_format.space_after = Pt(1)
        self.add_line_breaks(table_cell_1, 1)
        table_cell_1.add_paragraph(f"{reporting_team} conducted a {friendly_test_type_name} for {report_for} from {start_date} to {end_date}. The purpose of this assessment was to evaluate, and test risks associated with {
            report_for}'s infrastructure by simulating real-world attackers and identifying vulnerabilities.", style='rt_text_black')
        self.add_line_breaks(table_cell_1, 1)
        table_cell_1.add_paragraph('Client Contact:', style='Heading 3')
        table_cell_1.add_paragraph(
            "\n\n".join(client_contact_data),
            style='rt_text_black')
        self.add_line_breaks(table_cell_1, 1)
        table_cell_1.add_paragraph('Operators:', style='Heading 3')
        table_cell_1.add_paragraph(
            "\n\n".join(rt_contact_data),
            style='rt_text_black')
        for row in table.rows:
            row.height = Inches(10)
        # END OF page 3 Assessment Info

    def create_engagement_overview(
        self,
        reporting_team,
        report_for,
        test_type
    ):
        self.document.add_heading('Engagement Overview', level=1)
        self.document.add_paragraph(f"""{reporting_team} offers a comprehensive suite of security services, including penetration testing, adversary simulations, vulnerability assessments, external exposure evaluations, and full-scale red team engagements. Our manual, in-depth assessments go beyond automated scans to expose genuine attack vectors and strengthen your organization's defenses against sophisticated adversaries.  By emulating the tactics of advanced adversaries, {
            reporting_team} helps you identify gaps that traditional assessments might overlook. With decades of combined experience, {reporting_team} excels in application security, cloud security, Active Directory exploitation, and advanced red teaming. Our team of subject matter experts includes authorities in their fields who apply proven offensive techniques to enhance your security posture.""", style='rt_text_black')
        self.document.add_heading('Engagement Objectives', level=2)
        if test_type == "vuln_scan":
            self.document.add_paragraph(f"""This vulnerability assessment is designed to evaluate the security posture of {report_for} by conducting a comprehensive automated scan. The assessment focuses on identifying vulnerabilities that pose the most significant risk, including those that are most likely to be exploited and those that could result in the most critical impacts. By systematically analyzing {
                report_for}'s infrastructure, as defined in the scope below, this process provides valuable insights into potential attack vectors and areas requiring remediation to enhance overall security.""", style='rt_text_black')
        else:
            self.document.add_paragraph(f"""This penetration test is a targeted security assessment aimed at evaluating the resilience of {report_for} against real-world attack scenarios. Through a combination of manual and automated techniques, this test identifies exploitable vulnerabilities and assesses the potential impact of successful attacks. By simulating the actions of malicious actors, the penetration test provides an in-depth analysis of {
                report_for}'s infrastructure, applications, and defenses, as defined in the scope below. The findings will deliver actionable insights to help strengthen security measures and reduce the risk of compromise.""", style='rt_text_black')
        self.document.add_page_break()
        # END Engagement Overview page 4

    def create_attack_surface_overview(
        self,
        reporting_team,
    ):
        # START page 4
        self.document.add_heading('Engagement Overview', level=1)
        self.document.add_paragraph(f"""{reporting_team} offers a comprehensive suite of security services, including penetration testing, adversary simulations, vulnerability assessments, external exposure evaluations, and full-scale red team engagements. Our manual, in-depth assessments go beyond automated scans to expose genuine attack vectors and strengthen your organization's defenses against sophisticated adversaries.  By emulating the tactics of advanced adversaries, {
            reporting_team} helps you identify gaps that traditional assessments might overlook. With decades of combined experience, {reporting_team} excels in application security, cloud security, Active Directory exploitation, and advanced red teaming. Our team of subject matter experts includes authorities in their fields who apply proven offensive techniques to enhance your security posture.""", style='rt_text_black')
        self.document.add_heading('Service Description', level=2)
        self.document.add_paragraph(
            """Attack surface refers to all the potential avenues through which an attacker can enter a system, network, or application. It encompasses not only the software and hardware components but also the various protocols, interfaces, and interactions that can be targeted. Essentially, it represents the sum total of all points where an attacker could attempt to gain unauthorized access, disrupt operations, or extract information. Reducing the attack surface is a key aspect of enhancing cybersecurity, as it limits the number of potential entry points for attackers, thus making it harder for them to compromise a system.""",
            style='rt_text_black'
        )
        self.document.add_page_break()

    def create_process_methodology(
        self,
        reporting_team,
        report_for,
        special_considerations,
        test_type
    ):
        # START page 5 Process/Methodology
        if test_type != "vuln_scan":
            self.document.add_heading('Process and Methodology', level=1)
            self.document.add_paragraph(f"""{reporting_team} employed a thorough methodology to conduct a security assessment of {
                report_for}'s infrastructure, as outlined in the scope below. The process began with an in-depth analysis of the architecture and environment, including detailed scanning and automated testing to identify known vulnerabilities. This was followed by manual reconnaissance and customized exploitation to uncover and assess potential weaknesses.""", style='rt_text_black')
            if special_considerations:
                self.document.add_paragraph(
                    f"""{special_considerations}""",
                    style='rt_text_black')
            self.document.add_paragraph(
                """Reconnaissance""", style='List Number')
            recon_paragraph = self.document.add_paragraph(
                """The primary objective of the reconnaissance phase is to identify and verify all reachable endpoints within the defined scope. From these endpoints, ports, services, and protocols are enumerated to establish a detailed understanding of the environment. This foundational information enables a tailored and precise security assessment. Reconnaissance was conducted using a combination of automated scans and manual service fingerprinting and analysis. Additionally, Open Source Intelligence (OSINT) techniques were employed to gather relevant information about the environment prior to engaging directly with the infrastructure.""",
                style='rt_text_black')
            recon_paragraph.paragraph_format.left_indent = Inches(0.25)
            self.document.add_paragraph(
                """Exploration and Verification""",
                style='List Number')
            exploit_verification_paragraph = self.document.add_paragraph(f"""Using the findings from the reconnaissance phase, combined with their expertise and extensive experience, {reporting_team}'s assessors performed a thorough manual security analysis of {
                report_for}'s infrastructure, as defined within the scope below. The outcomes of this detailed exploration are documented in the assessment narrative and the vulnerability details section of this report.""", style='rt_text_black')
            exploit_verification_paragraph.paragraph_format.left_indent = Inches(
                0.25)
        self.document.add_paragraph(
            """Assessment Reporting""",
            style='List Number')
        reporting_paragraph = self.document.add_paragraph(
            f"""At the conclusion of the engagement, {reporting_team} provides {report_for}'s team with a comprehensive analysis and detailed report, including actionable recommendations. Our assessors uphold industry standards by delivering clear, concise reports that prioritize addressing the highest-risk vulnerabilities first.""",
            style='rt_text_black')
        reporting_paragraph.paragraph_format.left_indent = Inches(0.25)
        self.document.add_page_break()
        # END page 5 Process/Methodology

    def create_scope_and_rules(
        self,
        reporting_team,
        scope_external,
        scope_internal,
        wifi_scope,
        friendly_scope_test_from,
        ip_strip
    ):
        # START Scoping and rules of engagement page 6
        if not ip_strip:
            self.document.add_heading(
                'Scoping and Rules of Engagement', level=1)
            self.document.add_paragraph(f"""Although malicious actors operate without constraints, {
                reporting_team} recognizes the importance of defining a clear scope to ensure the assessment is completed efficiently and third parties not involved in the engagement are protected. The following limitations were applied to this engagement:""", style='rt_text_black')
            if scope_internal:
                scope_internal_paragraph = self.document.add_paragraph(
                    "Internal: ", style='List Bullet 2')
                scope_internal_paragraph.paragraph_format.left_indent = Inches(1)
                for int_scope in scope_internal:
                    tmp_int_scope = self.document.add_paragraph(
                        int_scope, style='List Bullet 2')
                    tmp_int_scope.paragraph_format.left_indent = Inches(1.5)
            if scope_external:
                scope_external_paragraph = self.document.add_paragraph(
                    "External: ", style='List Bullet 2')
                scope_external_paragraph.paragraph_format.left_indent = Inches(1)
                for ext_scope in scope_external:
                    tmp_ext_scope = self.document.add_paragraph(
                        ext_scope, style='List Bullet 2')
                    tmp_ext_scope.paragraph_format.left_indent = Inches(1.5)
            if wifi_scope:
                scope_wifi_paragraph = self.document.add_paragraph(
                    "WIFI: ", style='List Bullet 2')
                scope_wifi_paragraph.paragraph_format.left_indent = Inches(1)
                for wifi_scope in wifi_scope:
                    tmp_wifi_scope = self.document.add_paragraph(
                        wifi_scope, style='List Bullet 2')
                    tmp_wifi_scope.paragraph_format.left_indent = Inches(1.5)
            self.document.add_paragraph(f"""All assessment activities were conducted from {
                friendly_scope_test_from}, focusing exclusively on the scope items defined above. No testing was performed on any IP addresses, domains, or URLs outside the specified scope.""", style='List Bullet 2')
            self.document.add_page_break()
            # END rules of engagement page 6

    def create_risk_table(
        self,
        risk_level,
        risk_matrix_row,
        risk_matrix_col,
        risk_level_color_rgb,
        risk_level_color_hex
    ):
        risk_header = self.document.add_heading(level=3)
        risk_header.add_run('Overall Risk Rating: ')
        risk_title_color_block = risk_header.add_run("    ")
        risk_title_color_block.font.color.rgb = risk_level_color_rgb
        risk_title_color_block_tag = risk_title_color_block._r
        risk_title_color_block_shd = OxmlElement('w:shd')
        risk_title_color_block_shd.set(qn('w:val'), 'clear')
        risk_title_color_block_shd.set(qn('w:color'), 'auto')
        risk_title_color_block_shd.set(qn('w:fill'), risk_level_color_hex)
        risk_title_rpr = risk_title_color_block_tag.get_or_add_rPr()
        risk_title_rpr.append(risk_title_color_block_shd)
        risk_header.add_run(" {}".format(risk_level.title()))
        matrix_row_num = 6
        matrix_col_num = 7
        matrix_table = self.document.add_table(
            rows=matrix_row_num,
            cols=matrix_col_num,
            style="Table Grid")
        for row_num in range(0, matrix_row_num):
            for col_num in range(0, matrix_col_num):
                tmp_cell = matrix_table.cell(row_num, col_num)
                if (risk_matrix_row -
                    1) == row_num and (risk_matrix_col +
                                       1) == col_num:
                    self.set_cell_borders_black(tmp_cell)
                else:
                    self.set_cell_borders_clear_1(tmp_cell)
        matrix_table.autofit = False
        matrix_table.allow_autofit = False
        matrix_table_cell_00 = matrix_table.cell(0, 0)
        matrix_table_cell_00.merge(matrix_table.cell(4, 0))
        matrix_table_cell_00.width = Inches(1.3)
        matrix_table_cell_00.height = Inches(5)
        matrix_table_cell_00.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        matrix_table_cell_1 = matrix_table.cell(0, 1)
        matrix_table_cell_1.merge(matrix_table.cell(4, 1))
        matrix_table_cell_1.width = Inches(1.2)
        matrix_table_cell_1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        matrix_table_bottom_cell = matrix_table.cell(5, 2)
        matrix_table_bottom_cell.merge(matrix_table.cell(5, 6))
        matrix_table_bottom_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        matrix_key_info = matrix_table_cell_00.add_paragraph()
        matrix_key_label = matrix_key_info.add_run("KEY\n")
        info_color = matrix_key_info.add_run("    ")
        info_color.font.color.rgb = self.RISK_LEVEL_INFO_RGB
        info_tag = info_color._r
        info_shd = OxmlElement('w:shd')
        info_shd.set(qn('w:val'), 'clear')
        info_shd.set(qn('w:color'), 'auto')
        info_shd.set(qn('w:fill'), self.RISK_LEVEL_INFO_HEX)
        info_tag.rPr.append(info_shd)
        matrix_key_info.add_run(" Very Low\n")
        low_color = matrix_key_info.add_run("    ")
        low_color.font.color.rgb = self.RISK_LEVEL_LOW_RGB
        low_tag = low_color._r
        low_shd = OxmlElement('w:shd')
        low_shd.set(qn('w:val'), 'clear')
        low_shd.set(qn('w:color'), 'auto')
        low_shd.set(qn('w:fill'), self.RISK_LEVEL_LOW_HEX)
        low_tag.rPr.append(low_shd)
        matrix_key_info.add_run(" Low\n")
        moderate_color = matrix_key_info.add_run("    ")
        moderate_color.font.color.rgb = self.RISK_LEVEL_MODERATE_RGB
        moderate_tag = moderate_color._r
        moderate_shd = OxmlElement('w:shd')
        moderate_shd.set(qn('w:val'), 'clear')
        moderate_shd.set(qn('w:color'), 'auto')
        moderate_shd.set(qn('w:fill'), self.RISK_LEVEL_MODERATE_HEX)
        moderate_tag.rPr.append(moderate_shd)
        matrix_key_info.add_run(" Moderate\n")
        high_color = matrix_key_info.add_run("    ")
        high_color.font.color.rgb = self.RISK_LEVEL_HIGH_RGB
        high_tag = high_color._r
        high_shd = OxmlElement('w:shd')
        high_shd.set(qn('w:val'), 'clear')
        high_shd.set(qn('w:color'), 'auto')
        high_shd.set(qn('w:fill'), self.RISK_LEVEL_HIGH_HEX)
        high_tag.rPr.append(high_shd)
        matrix_key_info.add_run(" High\n")
        critical_color = matrix_key_info.add_run("    ")
        critical_color.font.color.rgb = self.RISK_LEVEL_CRITICAL_RGB
        critical_tag = critical_color._r
        critical_shd = OxmlElement('w:shd')
        critical_shd.set(qn('w:val'), 'clear')
        critical_shd.set(qn('w:color'), 'auto')
        critical_shd.set(qn('w:fill'), self.RISK_LEVEL_CRITICAL_HEX)
        critical_tag.rPr.append(critical_shd)
        matrix_key_info.add_run(" Critical")
        matrix_key_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        matrix_y_label = matrix_table_cell_1.add_paragraph(
            "Exploitation Likelihood")
        matrix_y_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        matrix_x_label = matrix_table_bottom_cell.add_paragraph(
            "Potential Impact")
        matrix_x_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in matrix_table.rows:
            row.height = Inches(0.6)
        info_cells = [
            matrix_table.cell(3, 2),
            matrix_table.cell(4, 2),
            matrix_table.cell(4, 3)
        ]
        for matrix_cell in info_cells:
            shading_elm_1 = parse_xml(
                r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), self.RISK_LEVEL_INFO_HEX))
            matrix_cell._tc.get_or_add_tcPr().append(shading_elm_1)
        low_cells = [
            matrix_table.cell(2, 2),
            matrix_table.cell(3, 3),
            matrix_table.cell(4, 4)
        ]
        for matrix_cell in low_cells:
            shading_elm_1 = parse_xml(
                r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), self.RISK_LEVEL_LOW_HEX))
            matrix_cell._tc.get_or_add_tcPr().append(shading_elm_1)
        moderate_cells = [
            matrix_table.cell(0, 2),
            matrix_table.cell(1, 2),
            matrix_table.cell(1, 3),
            matrix_table.cell(2, 3),
            matrix_table.cell(2, 4),
            matrix_table.cell(3, 4),
            matrix_table.cell(3, 5),
            matrix_table.cell(4, 5),
            matrix_table.cell(4, 6)
        ]
        for matrix_cell in moderate_cells:
            shading_elm_1 = parse_xml(
                r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), self.RISK_LEVEL_MODERATE_HEX))
            matrix_cell._tc.get_or_add_tcPr().append(shading_elm_1)
        high_cells = [
            matrix_table.cell(0, 3),
            matrix_table.cell(0, 4),
            matrix_table.cell(1, 4),
            matrix_table.cell(1, 5),
            matrix_table.cell(2, 5),
            matrix_table.cell(2, 6),
            matrix_table.cell(3, 6)
        ]
        for matrix_cell in high_cells:
            shading_elm_1 = parse_xml(
                r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), self.RISK_LEVEL_HIGH_HEX))
            matrix_cell._tc.get_or_add_tcPr().append(shading_elm_1)
        critical_cells = [
            matrix_table.cell(0, 5),
            matrix_table.cell(0, 6),
            matrix_table.cell(1, 6)
        ]
        for matrix_cell in critical_cells:
            shading_elm_1 = parse_xml(
                r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), self.RISK_LEVEL_CRITICAL_HEX))
            matrix_cell._tc.get_or_add_tcPr().append(shading_elm_1)

    def create_executive_summary(
        self,
        reporting_team,
        report_for,
        risk_level,
        test_type,
        friendly_test_type_name,
        friendly_compromise_status,
        risk_level_color_rgb,
        risk_level_color_hex,
        risk_matrix_row,
        risk_matrix_col
    ):
        # START page 7 Exec Summary
        self.document.add_heading('Executive summary', level=1)
        self.document.add_paragraph(f"""{reporting_team} conducted a {friendly_test_type_name} to evaluate the security posture of {
            report_for}'s infrastructure, as defined within the outlined scope. This assessment aimed to proactively identify vulnerabilities, validate their severity, and provide actionable recommendations to enhance {report_for}'s overall security.""", style='rt_text_black')

        if test_type != "vuln_scan":
            self.document.add_paragraph(f"""The assessment revealed a {risk_level.lower()} risk of compromise within {report_for}'s infrastructure, based on the findings detailed in this report. {reporting_team}'s assessors were {
                friendly_compromise_status} in exploiting in-scope systems and leveraging identified misconfigurations. Key areas for improvement have been highlighted, with detailed findings and remediation recommendations provided in subsequent sections of this report.""", style='rt_text_black')

        self.document.add_heading('{} Risk Rating'.format(report_for), level=2)
        self.document.add_paragraph(f"""The calculated risk level of {risk_level.lower()} is based on the likelihood of exploitation (ease of exploitation) and the potential business impact on {
            report_for}'s environment. Proactively addressing these findings will significantly strengthen {report_for}'s defenses against potential threats.""", style='rt_text_black')
        # create risk title and risk table
        _ = self.create_risk_table(
            risk_level,
            risk_matrix_row,
            risk_matrix_col,
            risk_level_color_rgb,
            risk_level_color_hex
        )
        self.document.add_page_break()
        # END exec summary page 7

    def create_findings_overviews(
        self,
        vulnerability_dict
    ):
        # START findings overviews
        self.document.add_heading('Findings Overview', level=1)
        self.document.add_paragraph(
            "Risk level is determined by evaluating both the likelihood of exploitation and the potential impact on the organization. The risk decreases when an adversary is required to expend significant resources to achieve their objectives. The greater the effort and resources needed to gain access, the lower the likelihood of exploitation. Additionally, when attackers encounter a strong security posture, they are more inclined to shift their focus to less secure targets, reducing the threat to your organization.",
            style='rt_text_black')
        vulnerability_count = sum([len(risk_dict.keys())
                                   for _, risk_dict in vulnerability_dict.items()])
        findings_overview_table = self.document.add_table(
            vulnerability_count + 1, 2)
        finding_id_cell = findings_overview_table.cell(0, 0)
        self.set_cell_borders_black_bottom(finding_id_cell)
        finding_id_cell.width = Inches(7)
        finding_id_paragraph = finding_id_cell.add_paragraph(
            "Finding ID - Description", style="findings_table_text_title")
        finding_risk_level_cell = findings_overview_table.cell(0, 1)
        self.set_cell_borders_black_bottom(finding_risk_level_cell)
        finding_risk_level_paragraph = finding_risk_level_cell.add_paragraph(
            "Risk Level", style="findings_table_text_title")
        finding_risk_level_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        row_num = 1
        findings_sort_order = (
            'critical',
            'high',
            'medium',
            'moderate',
            'low',
            'info',
            'informational'
        )
        ordered_findings_list = []
        for finding_level in findings_sort_order:
            ordered_findings_list.append(
                (finding_level, vulnerability_dict.get(finding_level)))
        for risk_level, risk_dict in ordered_findings_list:
            if risk_level.lower() == "medium":
                risk_level = "moderate"
            if risk_level.lower() == "info":
                risk_level = "informational"
            tmp_risk_num = 1
            if risk_dict:
                for risk_name, risk_data in risk_dict.items():
                    risk_name_start = "{}{}".format(
                        risk_level[0].upper(), tmp_risk_num)
                    finding_title = '{} -  {}'.format(
                        risk_name_start, risk_name)
                    tmp_finding_title_cell = findings_overview_table.cell(
                        row_num, 0)
                    tmp_finding_title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    self.set_findings_table_margins(tmp_finding_title_cell)
                    tmp_finding_title_paragraph = tmp_finding_title_cell.paragraphs[0]
                    tmp_finding_title_paragraph.add_run(
                        finding_title, style="findings_table_text")
                    tmp_finding_title_paragraph.paragraph_format.space_after = Pt(
                        0)
                    tmp_finding_title_paragraph.paragraph_format.space_before = Pt(
                        0)
                    tmp_risk_level_cell = findings_overview_table.cell(
                        row_num, 1)
                    self.set_findings_table_margins(tmp_risk_level_cell)
                    tmp_risk_level_paragraph = tmp_risk_level_cell.paragraphs[0]
                    tmp_risk_level_paragraph_run = tmp_risk_level_paragraph.add_run(
                        risk_level.title(), style="findings_table_text_alt")
                    tmp_risk_level_paragraph_run.font.color.rgb = self.risk_level_color(risk_level)[
                        0]
                    tmp_risk_level_paragraph.paragraph_format.space_after = Pt(
                        0)
                    tmp_risk_level_paragraph.paragraph_format.space_before = Pt(
                        0)
                    tmp_risk_level_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    tmp_risk_level_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tmp_risk_num += 1
                    row_num += 1
        self.document.add_page_break()
        # END findings overviews
        return ordered_findings_list

    def create_findings_and_recommendations(
        self,
        ordered_findings_list,
        ip_strip,
        with_evidence
    ):
        # START findings and recommendations
        self.document.add_heading('Findings and Recommendations', level=1)
        # FOR LOOP THIS TO MAKE FINDINGS
        risk_data_list_count = sum(
            [len(risk_dict.keys()) for _, risk_dict in ordered_findings_list if risk_dict])
        risk_iteration_count = 1
        for risk_level, risk_dict in ordered_findings_list:
            tmp_risk_num = 1
            if risk_dict:
                for risk_name, risk_data in risk_dict.items():
                    if risk_data:
                        risk_name_start = "{}{}".format(
                            risk_level[0].upper(), tmp_risk_num)
                        self.document.add_heading('{} –  {}'.format(
                            risk_name_start, risk_name), level=2)
                        self.document.add_paragraph(
                            "_______________________________________________________________________________________________")
                        risk_level_finding_level = self.document.add_paragraph()
                        risk_level_finding_title = risk_level_finding_level.add_run(
                            "Risk Level: ")
                        risk_level_finding_level_label = risk_level_finding_level.add_run(
                            "{}".format(risk_level.upper()))
                        risk_level_finding_level_label.font.color.rgb = self.risk_level_color(risk_level)[
                            0]
                        self.document.add_paragraph(
                            "_______________________________________________________________________________________________")
                        self.document.add_paragraph(
                            "Description", style="rt_header_text_black")
                        self.document.add_paragraph(
                            risk_data.get('description'), style="rt_text_black")
                        screenshot_list = risk_data.get('finding_screenshots')
                        if screenshot_list:
                            for screenshot_dict in screenshot_list:
                                screenshot_paragraph = self.document.add_paragraph(
                                    style="rt_text_center_black_subtext")
                                screenshot_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                screenshot_centered = screenshot_paragraph.add_run()
                                # add pic
                                screenshot_pic = screenshot_centered.add_picture(
                                    screenshot_dict.get('finding_screenshot'), width=Inches(5))
                                # add centered desc
                                screenshot_subtile = self.document.add_paragraph(
                                    screenshot_dict.get('finding_subtle'), style="rt_text_center_black_subtext")
                                screenshot_subtile.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if not ip_strip:
                            if risk_data.get('hosts_list'):
                                self.document.add_paragraph(
                                    "Affected Hosts", style="rt_header_text_black")
                                for affected_host in list(
                                        set(risk_data.get('hosts_list'))):
                                    self.document.add_paragraph(
                                        "{}".format(affected_host), style='List Bullet 2')
                        self.document.add_paragraph(
                            "Impact", style="rt_header_text_black")
                        self.document.add_paragraph(
                            risk_data.get('impact'), style="rt_text_black")
                        self.document.add_paragraph(
                            "Recommendations", style="rt_header_text_black")
                        self.document.add_paragraph(
                            risk_data.get('recommendations'), style="rt_text_black")
                        # solutions screenshots here
                        solutions_screenshot_list = risk_data.get(
                            'recommendations_screenshots')
                        if solutions_screenshot_list:
                            for screenshot_dict in solutions_screenshot_list:
                                screenshot_paragraph = self.document.add_paragraph(
                                    style="rt_text_center_black_subtext")
                                screenshot_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                screenshot_centered = screenshot_paragraph.add_run()
                                # add pic
                                screenshot_pic = screenshot_centered.add_picture(
                                    screenshot_dict.get('finding_screenshot'), width=Inches(5))
                                # add centered desc
                                screenshot_subtile = self.document.add_paragraph(
                                    screenshot_dict.get('finding_subtle'), style="rt_text_center_black_subtext")
                                screenshot_subtile.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # add See Also
                        if risk_data.get('sources_list'):
                            self.document.add_paragraph(
                                "Sources", style="rt_header_text_black")
                            for source_item in list(
                                    set(risk_data.get('sources_list'))):
                                self.document.add_paragraph(
                                    "{}".format(source_item), style='List Bullet 2')
                        # add CVEs
                        if risk_data.get('cve_list'):
                            self.document.add_paragraph(
                                "CVEs", style="rt_header_text_black")
                            for cve_number in list(
                                    set(risk_data.get('cve_list'))):
                                self.document.add_paragraph(
                                    "{}".format(cve_number), style='List Bullet 2')
                        if with_evidence:
                            if not ip_strip:
                                # add evidence
                                if risk_data.get('evidence'):
                                    self.document.add_paragraph(
                                        "Evidence", style="rt_header_text_black")
                                    for evidence_item in list(
                                            set(risk_data.get('evidence'))):
                                        if evidence_item:
                                            self.document.add_paragraph(
                                                "{}\n\n".format(evidence_item), style='rt_text_black')
                        tmp_risk_num += 1
                        if risk_iteration_count < risk_data_list_count:
                            self.document.add_page_break()
                        risk_iteration_count += 1
        # END findings and recommendations

    def create_attack_surface_findings(
        self,
        port_protocol_desc_list,
        plain_text_port_list,
        screenshot_http_file_list,
        ssl_cipher_list,
        reject_protocol_desc_list
    ):
        # START findings and recommendations
        self.document.add_heading('Findings and Recommendations', level=1)
        # FOR LOOP THIS TO MAKE FINDINGS
        if port_protocol_desc_list:
            port_list = list(set(
                [d.split(":")[3] for d in port_protocol_desc_list if d.split(":")[3] != "port"]))
            if [port for port in port_list if port in plain_text_port_list] or len(
                    screenshot_http_file_list) > 10:
                risk_level = "critical"
            elif 20 > len(port_list):
                risk_level = "high"
            elif 19 > len(port_list) > 10:
                risk_level = "medium"
            else:
                risk_level = "low"
            risk_name = "Exposed Ports"
            self.document.add_heading(risk_name, level=2)
            self.document.add_paragraph(
                "_______________________________________________________________________________________________")
            risk_level_finding_level = self.document.add_paragraph()
            risk_level_finding_title = risk_level_finding_level.add_run(
                "Risk Level: ")
            risk_level_finding_level_label = risk_level_finding_level.add_run(
                "{}".format(risk_level.upper()))
            risk_level_finding_level_label.font.color.rgb = self.risk_level_color(risk_level)[
                0]
            self.document.add_paragraph(
                "_______________________________________________________________________________________________")
            self.document.add_paragraph(
                "Description", style="rt_header_text_black")
            description = """Exposing services to the internet means making them accessible from anywhere on the web. While this might seem convenient for users or clients, it opens up your systems to potential threats from malicious actors, including hackers, malware, and automated bots. Even seemingly harmless services can become entry points for attacks if left unprotected."""
            self.document.add_paragraph(description, style="rt_text_black")
            if screenshot_http_file_list:
                for screenshot_filename in screenshot_http_file_list:
                    screenshot_paragraph = self.document.add_paragraph(
                        style="rt_text_center_black_subtext")
                    screenshot_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    screenshot_centered = screenshot_paragraph.add_run()
                    # add pic
                    screenshot_pic = screenshot_centered.add_picture(
                        screenshot_filename, width=Inches(6))
                    # add centered desc
                    screenshot_subtile = self.document.add_paragraph(
                        screenshot_filename.split("/")[-1], style="rt_text_center_black_subtext")
                    screenshot_subtile.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.document.add_paragraph(
                "Affected Hosts", style="rt_header_text_black")
            for affected_host in port_protocol_desc_list:
                self.document.add_paragraph(
                    affected_host, style='List Bullet 2')
            self.document.add_paragraph(
                "Impact", style="rt_header_text_black")
            impact = """When services are exposed to the internet, they become potential targets for malicious actors seeking to exploit vulnerabilities for unauthorized access or nefarious purposes. Without adequate authentication, authorization mechanisms, and protections anyone could potentially breach the system, leading to data breaches or misuse of resources. Additionally, cyberattacks like DDoS attacks, SQL injections, and malware infections pose significant threats, compromising the availability, confidentiality, and integrity of sensitive information. Furthermore, compliance and legal issues may arise due to data privacy regulations, necessitating stringent security measures to avoid penalties and reputational damage."""
            self.document.add_paragraph(impact, style="rt_text_black")
            self.document.add_paragraph(
                "Recommendations", style="rt_header_text_black")
            recommendations = """To mitigate the risks associated with exposing services to the internet, organizations should prioritize security measures. This includes implementing robust authentication and authorization protocols, utilizing encryption for data transmission, regularly updating and patching systems to address vulnerabilities, and establishing comprehensive monitoring and logging mechanisms to detect and respond to suspicious activities promptly. Additionally, limiting the exposure of services to the internet and considering alternative solutions such as Zero Trust Tunnels, VPNs, or private networks can further enhance security. By adopting a security-first approach, organizations can safeguard their systems and data from potential threats while ensuring compliance with regulatory requirements. """
            self.document.add_paragraph(recommendations, style="rt_text_black")
            self.document.add_page_break()
        if ssl_cipher_list:
            risk_level = "high"
            risk_name = "Legacy Ciphers/Certificates"
            self.document.add_heading(risk_name, level=2)
            self.document.add_paragraph(
                "_______________________________________________________________________________________________")
            risk_level_finding_level = self.document.add_paragraph()
            risk_level_finding_title = risk_level_finding_level.add_run(
                "Risk Level: ")
            risk_level_finding_level_label = risk_level_finding_level.add_run(
                "{}".format(risk_level.upper()))
            risk_level_finding_level_label.font.color.rgb = self.risk_level_color(risk_level)[
                0]
            self.document.add_paragraph(
                "_______________________________________________________________________________________________")
            self.document.add_paragraph(
                "Description", style="rt_header_text_black")
            description = """Harden SSL/TLS ciphers and certificates is essential for enhancing the security of communication over the internet. SSL/TLS protocols are fundamental for encrypting data transmissions between clients and servers, ensuring confidentiality and integrity. However, outdated or weak ciphers and certificates can leave systems vulnerable to various attacks, compromising the security of sensitive information. Therefore, hardening SSL/TLS configurations is crucial to mitigate risks and maintain robust security standards."""
            self.document.add_paragraph(description, style="rt_text_black")
            self.document.add_paragraph(
                "Affected Hosts", style="rt_header_text_black")
            for affected_host in ssl_cipher_list:
                self.document.add_paragraph(
                    affected_host, style='List Bullet 2')
            self.document.add_paragraph("Impact", style="rt_header_text_black")
            impact = """SSL/TLS ciphers and certificates play a critical role in establishing secure connections between clients and servers by encrypting data transmissions. However, over time, vulnerabilities are discovered in encryption algorithms and certificate authorities, rendering certain ciphers and certificates susceptible to exploitation by malicious actors. Weak ciphers can be exploited to intercept or tamper with encrypted data, while compromised or improperly configured certificates can lead to man-in-the-middle attacks or unauthorized access to sensitive information. Therefore, it's imperative to regularly assess and update SSL/TLS configurations to ensure they adhere to the latest security standards and best practices."""
            self.document.add_paragraph(impact, style="rt_text_black")
            self.document.add_paragraph(
                "Recommendations", style="rt_header_text_black")
            recommendations = """To harden SSL/TLS ciphers and certificates, organizations should adopt several best practices. This includes disabling outdated or vulnerable ciphers and protocols, such as SSLv2, SSLv3, TLSv1, and TLSv1.1, and prioritizing the use of strong, modern encryption algorithms like AES and ChaCha20. Additionally, regularly updating SSL/TLS libraries and software to patch known vulnerabilities is essential for maintaining security. Furthermore, organizations should implement proper certificate management practices, including regularly renewing certificates, using trusted certificate authorities, and deploying mechanisms like certificate pinning to prevent unauthorized certificate issuance. By proactively hardening SSL/TLS configurations, organizations can bolster the security of their communication channels and protect sensitive data from potential threats."""
            self.document.add_paragraph(recommendations, style="rt_text_black")
            self.document.add_paragraph(
                "Sources", style="rt_header_text_black")
            sources_list = [
                "https://www.kali.org/tools/sslscan/",
                "https://www.ssllabs.com/ssltest/",
                "https://ssl-config.mozilla.org/",
                "https://wiki.mozilla.org/Security/Server_Side_TLS",
            ]
            for source_item in sources_list:
                self.document.add_paragraph(
                    f"{source_item}", style='List Bullet 2')
            self.document.add_page_break()
        if reject_protocol_desc_list:
            risk_level = "low"
            risk_name = "Firewall Rule Set to RST/REJECT Instead of DROP"
            self.document.add_heading(risk_name, level=2)
            self.document.add_paragraph("_"*95)
            risk_level_finding_level = self.document.add_paragraph()
            risk_level_finding_title = risk_level_finding_level.add_run(
                "Risk Level: ")
            risk_level_finding_level_label = risk_level_finding_level.add_run(
                "{}".format(risk_level.upper()))
            risk_level_finding_level_label.font.color.rgb = self.risk_level_color(risk_level)[0]
            self.document.add_paragraph("_"*95)
            self.document.add_paragraph(
                "Description", style="rt_header_text_black")
            description = """Choosing between DROP and REJECT rules on a firewall entails considering the trade-offs between stealth and responsiveness. DROP rules silently discard packets without notifying the sender, while REJECT rules send back an explicit rejection(reset packet) message. The decision hinges on the desired balance between security and network visibility, as DROP rules can obscure the existence of protected resources, making them less susceptible to reconnaissance, but may hinder troubleshooting efforts. Therefore, understanding the implications of each approach is crucial for effective firewall management and network security."""
            self.document.add_paragraph(description, style="rt_text_black")
            self.document.add_paragraph(
                "Affected Hosts", style="rt_header_text_black")
            for affected_host in reject_protocol_desc_list:
                self.document.add_paragraph(
                    affected_host, style='List Bullet 2')
            self.document.add_paragraph("Impact", style="rt_header_text_black")
            impact = """DROP rules on a firewall provide a higher level of stealth and security by silently discarding packets that match the specified criteria. This means that attackers attempting to probe the network are less likely to receive any indication that their probes have reached a protected resource, making it harder for them to gather information about potential targets. On the other hand, REJECT rules explicitly notify the sender that their connection attempt has been rejected(reset packet), which can inadvertently provide attackers with valuable reconnaissance information about the network topology and potentially aid in crafting more targeted attacks."""
            self.document.add_paragraph(impact, style="rt_text_black")
            self.document.add_paragraph(
                "Recommendations", style="rt_header_text_black")
            recommendations = """When configuring firewall rules, organizations should carefully consider their security requirements and operational needs. In high-security environments where stealth is paramount, such as protecting critical infrastructure or sensitive data, prioritizing DROP rules may be advisable to minimize the exposure of protected resources to potential attackers."""
            self.document.add_paragraph(recommendations, style="rt_text_black")
            self.document.add_page_break()

    def save_the_doc(self,
            ip_strip,
            test_type,
            report_for
        ):
        # SAVE THE DOC
        if ip_strip:
            report_name = report_for.replace(' ', '_').lower()
            file_name_saved = f"{self.save_dir}{
                report_name}_exec_{test_type}.docx"
        else:
            report_name = report_for.replace(' ', '_').lower()
            file_name_saved = f"{self.save_dir}{report_name}_{test_type}.docx"
        self.document.save(file_name_saved)
        return file_name_saved

    def generate_doc(
        self,
        REPORT_FOR="Company Name",
        TITLE_MONTH_YEAR="October 2021",
        TEST_TYPE="pentest",
        START_DATE="01 October 2021",
        END_DATE="31 October 2021",
        CLIENT_CONTACT_DATA=[
            "First Name\nSenior Analyst\nCompany Name\n(555) 123-4567\nfanme@company.com",
        ],
        RT_CONTACT_DATA=[
            "Leon Denard\nRed Team Lead\nldenard@redteam-ioc-test.com"
        ],
        SCOPE_INTERNAL=["10.0.0.0/8"],
        SCOPE_EXTERNAL=["domain.com", "1.1.1.1/32"],
        WIFI_SCOPE=[],
        SCOPE_TEST_FROM="both",
        COMPROMISE_SUCCESS=True,
        RISK_LEVEL="Critical",
        RISK_MATRIX_ROW=1,
        RISK_MATRIX_COL=4,
        VULN_SCAN_CSV_FILE_PATH="",
        VULN_LIST=[],
        filter_to_scores=False,
        ip_strip=False,
        score_overrides_dict={},
        strength_list=[],
        improvements_list=[],
        with_evidence=False,
        SPECIAL_CONSIDERATIONS='',
        solutions_override_dict={},
        REPORTING_TEAM="Red Team",
    ):
        FRIENDLY_TEST_TYPE_NAME = self.get_friendly_test_name(TEST_TYPE)
        FRIENDLY_SCOPE_TEST_FROM = self.get_friendly_scope_test_from(
            SCOPE_TEST_FROM)
        FRIENDLY_COMPROMISE_STATUS = self.get_compromise_status(
            COMPROMISE_SUCCESS)
        RISK_LEVEL_COLOR_RGB, RISK_LEVEL_COLOR_HEX = self.get_risk_color_codes(
            RISK_LEVEL)
        VULNERABILITY_DICT = self.format_vuln_data_into_dict(
            VULN_SCAN_CSV_FILE_PATH,
            VULN_LIST,
            filter_to_scores,
            score_overrides_dict,
            solutions_override_dict
        )
        # setup document and styles
        _ = self.setup_document_and_styles()
        # create title page
        _ = self.create_title_page(
            REPORTING_TEAM,
            REPORT_FOR,
            FRIENDLY_TEST_TYPE_NAME,
            TITLE_MONTH_YEAR
        )
        # create Table of Contents
        _ = self.create_table_of_contents()
        # create assessment contact page
        _ = self.create_assessment_contact(
            REPORTING_TEAM,
            REPORT_FOR,
            FRIENDLY_TEST_TYPE_NAME,
            START_DATE,
            END_DATE,
            RT_CONTACT_DATA,
            CLIENT_CONTACT_DATA
        )
        # create engagement overview page
        _ = self.create_engagement_overview(
            REPORTING_TEAM,
            REPORT_FOR,
            TEST_TYPE
        )
        # create process and methodology page
        _ = self.create_process_methodology(
            REPORTING_TEAM,
            REPORT_FOR,
            SPECIAL_CONSIDERATIONS,
            TEST_TYPE
        )
        # create scoping and rules of engagement page
        _ = self.create_scope_and_rules(
            REPORTING_TEAM,
            SCOPE_EXTERNAL,
            SCOPE_INTERNAL,
            WIFI_SCOPE,
            FRIENDLY_SCOPE_TEST_FROM,
            ip_strip
        )
        # create executive summary page
        _ = self.create_executive_summary(
            REPORTING_TEAM,
            REPORT_FOR,
            RISK_LEVEL,
            TEST_TYPE,
            FRIENDLY_TEST_TYPE_NAME,
            FRIENDLY_COMPROMISE_STATUS,
            RISK_LEVEL_COLOR_RGB,
            RISK_LEVEL_COLOR_HEX,
            RISK_MATRIX_ROW,
            RISK_MATRIX_COL
        )
        # create findings overviews page
        ordered_findings_list = self.create_findings_overviews(
            VULNERABILITY_DICT
        )
        # create findings and recommendation pages
        _ = self.create_findings_and_recommendations(
            ordered_findings_list,
            ip_strip,
            with_evidence
        )
        # save the file
        file_name_saved = self.save_the_doc(
            ip_strip,
            TEST_TYPE,
            REPORT_FOR
        )
        # set the Table of Contents trigger to update on open
        self.set_updatefields_true(file_name_saved)
        return file_name_saved

    def attack_surface_generate_doc(
        self,
        REPORT_FOR="Company Name",
        TITLE_MONTH_YEAR="March 2024",
        START_DATE="19 March 2024",
        END_DATE="19 March 2024",
        RT_CONTACT_DATA=[
            "Leon Denard\nRed Team Lead\nldenard@redteam-test.com"
        ],
        CLIENT_CONTACT_DATA=[],
        SCOPE_EXTERNAL=["domain.com", "1.1.1.1/32"],
        port_protocol_desc_list=[],
        reject_protocol_desc_list=[],
        screenshot_http_file_list=[],
        ssl_cipher_list=[],
        plain_text_port_list=[],
        REPORTING_TEAM="Red Team",
        SCOPE_TEST_FROM="external"
    ):
        FRIENDLY_SCOPE_TEST_FROM = self.get_friendly_scope_test_from(
            SCOPE_TEST_FROM)
        FRIENDLY_TEST_TYPE_NAME = "attack surface"
        # setup document and styles
        _ = self.setup_document_and_styles()
        # create title page
        _ = self.create_title_page(
            REPORTING_TEAM,
            REPORT_FOR,
            FRIENDLY_TEST_TYPE_NAME,
            TITLE_MONTH_YEAR
        )
        # create Table of Contents
        _ = self.create_table_of_contents()
        # create assessment contact page
        _ = self.create_assessment_contact(
            REPORTING_TEAM,
            REPORT_FOR,
            FRIENDLY_TEST_TYPE_NAME,
            START_DATE,
            END_DATE,
            RT_CONTACT_DATA,
            CLIENT_CONTACT_DATA
        )
        # create overview
        _ = self.create_attack_surface_overview(
            REPORTING_TEAM,
        )
        # create scoping page
        _ = self.create_scope_and_rules(
            REPORTING_TEAM,
            SCOPE_EXTERNAL,
            [],
            [],
            FRIENDLY_SCOPE_TEST_FROM,
            False
        )
        # generate the findings
        _ = self.create_attack_surface_findings(
            port_protocol_desc_list,
            plain_text_port_list,
            screenshot_http_file_list,
            ssl_cipher_list,
            reject_protocol_desc_list
        )
        # save the file
        file_name_saved = self.save_the_doc(
            False,
            "attack_surface",
            REPORT_FOR
        )
        # set the Table of Contents trigger to update on open
        self.set_updatefields_true(file_name_saved)
        return file_name_saved
